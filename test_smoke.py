"""Smoke tests for Scoped RAG modules.
Each check reports the owning file and symbol along with the outcome.
"""

from __future__ import annotations

import sys
from dataclasses import dataclass
from pathlib import Path
from typing import Callable, Tuple

# Ensure project root is importable
PROJECT_ROOT = Path(__file__).parent
sys.path.insert(0, str(PROJECT_ROOT))

CheckResult = Tuple[bool, str]


@dataclass
class CheckSpec:
    name: str
    file: str
    target: str
    runner: Callable[[], CheckResult]


def check_ocr_module() -> CheckResult:
    try:
        from ocr import (
            DoclingOCREngine,
            get_file_stats,
            get_supported_extensions,
            is_supported_file,
        )
        engine = DoclingOCREngine(enable_ocr=False, enable_tables=False, max_workers=1)
        _ = engine  # Silence linter warning for unused variable
        extensions = get_supported_extensions()
        stats = get_file_stats()
        probe = is_supported_file("example.pdf")
        detail = (
            f"extensions={len(extensions)} | example.pdf supported={probe} | "
            f"catalogued_files={stats.get('total', 0)}"
        )
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_image_helpers() -> CheckResult:
    try:
        from image import SUPPORTED_FORMATS, get_image_format
        fmt = get_image_format(Path("test.jpg"))
        detail = f"formats={sorted(SUPPORTED_FORMATS)} | .jpg-> {fmt}"
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_audio_metadata() -> CheckResult:
    try:
        from audio import extract_file_metadata
        sample = Path("sample_[L3].mp3")
        meta = extract_file_metadata(sample)
        detail = (
            f"access_level={meta.get('access_level')} | keywords={meta.get('keywords', '')}"
        )
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_vector_store() -> CheckResult:
    try:
        from langchain_core.documents import Document
        from vector_store import create_vector_store

        dummy = Document(page_content="test", metadata={"source": "test"})
        store = create_vector_store([dummy])
        ids = store._collection.count() if hasattr(store, "_collection") else "unknown"
        detail = f"created_store | chunks_stored={ids}"
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_splitter() -> CheckResult:
    try:
        from langchain_core.documents import Document
        from langchain_text_splitters import RecursiveCharacterTextSplitter

        splitter = RecursiveCharacterTextSplitter(chunk_size=64, chunk_overlap=16)
        doc = Document(page_content=("Sentence." * 10), metadata={"source": "unit"})
        chunks = splitter.split_documents([doc])
        detail = f"chunks_created={len(chunks)}"
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_docling_import() -> CheckResult:
    try:
        from docling.document_converter import DocumentConverter  # type: ignore
        from docling.datamodel.pipeline_options import PdfPipelineOptions  # type: ignore

        detail = (
            f"DocumentConverter={DocumentConverter.__name__} | "
            f"PdfPipelineOptions={PdfPipelineOptions.__name__}"
        )
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_office_fallbacks() -> CheckResult:
    missing = []
    try:
        from docx import Document as _  # type: ignore
    except Exception:
        missing.append("python-docx")

    try:
        from pptx import Presentation  # type: ignore
        _ = Presentation
    except Exception:
        missing.append("python-pptx")

    try:
        from openpyxl import load_workbook  # type: ignore
        _ = load_workbook
    except Exception:
        missing.append("openpyxl")

    try:
        from pypdf import PdfReader  # type: ignore
        _ = PdfReader
    except Exception:
        missing.append("pypdf")

    if missing:
        return False, f"missing={', '.join(missing)}"
    return True, "all optional fallbacks available"


def check_ollama_connection() -> CheckResult:
    try:
        import ollama

        models = ollama.list()
        names = [m.get("name") or m.get("model") for m in models.get("models", [])]
        detail = f"models_seen={len(names)}"
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_env_config() -> CheckResult:
    try:
        from config import os  # noqa: F401
        value = os.environ.get("GOOGLE_API_KEY")
        if value and "your_gemini_api_key_here" not in value:
            return True, "GOOGLE_API_KEY loaded from environment"
        return False, "GOOGLE_API_KEY missing or still placeholder"
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_groq_api_key() -> CheckResult:
    """Check if GROQ_API_KEY is configured"""
    try:
        import os
        from dotenv import load_dotenv
        load_dotenv()
        value = os.environ.get("GROQ_API_KEY")
        if value and len(value) > 10:
            return True, f"GROQ_API_KEY loaded (length={len(value)})"
        return False, "GROQ_API_KEY missing or invalid"
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_bm25_import() -> CheckResult:
    """Check BM25 library for hybrid search"""
    try:
        from rank_bm25 import BM25Okapi
        
        # Test basic functionality
        corpus = [
            "hello world".split(),
            "hello python".split(),
            "goodbye world".split(),
        ]
        bm25 = BM25Okapi(corpus)
        scores = bm25.get_scores("hello".split())
        
        detail = f"BM25Okapi available | test_scores={[round(s, 2) for s in scores]}"
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_hybrid_search() -> CheckResult:
    """Check hybrid search functionality"""
    try:
        from vector_store import (
            BM25Index,
            ENABLE_HYBRID_SEARCH,
            ENABLE_RERANKING,
            ENABLE_QUERY_EXPANSION,
        )
        from langchain_core.documents import Document
        
        # Test BM25Index
        index = BM25Index()
        docs = [
            Document(page_content="Python is a programming language", metadata={"access_level": 1}),
            Document(page_content="Java is also a programming language", metadata={"access_level": 2}),
            Document(page_content="Machine learning uses Python", metadata={"access_level": 3}),
        ]
        index.add_documents(docs)
        results = index.search("Python programming", k=2, access_level=5)
        
        detail = (
            f"BM25Index works | hybrid={ENABLE_HYBRID_SEARCH} | "
            f"rerank={ENABLE_RERANKING} | expand={ENABLE_QUERY_EXPANSION} | "
            f"search_results={len(results)}"
        )
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_query_expansion() -> CheckResult:
    """Check query expansion functionality"""
    try:
        from vector_store import expand_query, ENABLE_QUERY_EXPANSION
        
        test_query = "How does machine learning work?"
        queries = expand_query(test_query)
        
        if ENABLE_QUERY_EXPANSION:
            detail = f"expanded '{test_query[:30]}...' → {len(queries)} variations"
        else:
            detail = f"query expansion disabled | passthrough={len(queries)} query"
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_reranking() -> CheckResult:
    """Check reranking functionality"""
    try:
        from vector_store import rerank_documents, ENABLE_RERANKING, RERANK_TOP_K
        from langchain_core.documents import Document
        
        docs = [
            Document(page_content="Python is great for data science", metadata={}),
            Document(page_content="JavaScript is for web development", metadata={}),
            Document(page_content="Python machine learning tutorial", metadata={"heading_context": "Python ML Guide"}),
            Document(page_content="Random unrelated content here", metadata={}),
        ]
        
        query = "Python machine learning"
        reranked = rerank_documents(query, docs, top_k=3)
        
        detail = (
            f"reranking={'enabled' if ENABLE_RERANKING else 'disabled'} | "
            f"top_k={RERANK_TOP_K} | input={len(docs)} → output={len(reranked)}"
        )
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_llm_cascade() -> CheckResult:
    """Check LLM cascade configuration"""
    try:
        from config import (
            LLM_MODEL_PRIMARY,
            LLM_MODEL_SECONDARY,
            LLM_MODEL_TERTIARY,
            LLM_MODEL_LOCAL,
        )
        from llm import get_llm
        
        # Just check configuration, don't actually invoke
        detail = (
            f"primary={LLM_MODEL_PRIMARY.split('/')[-1][:20]} | "
            f"secondary={LLM_MODEL_SECONDARY.split('/')[-1][:15]} | "
            f"tertiary={LLM_MODEL_TERTIARY.split('/')[-1][:15]} | "
            f"local={LLM_MODEL_LOCAL.split(':')[0]}"
        )
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_llm_invocation() -> CheckResult:
    """Check if LLM can be invoked (tests cascade fallback)"""
    try:
        from llm import get_llm
        
        llm, model_name = get_llm()
        response = llm.invoke("Say 'test ok' in exactly 2 words")
        
        # Handle both string and AIMessage responses
        if hasattr(response, 'content'):
            answer = response.content
        else:
            answer = str(response)
        
        detail = f"model={model_name} | preview='{answer[:50]}...'" if len(answer) > 50 else f"model={model_name} | answer='{answer}'"
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


def check_embeddings() -> CheckResult:
    """Check embedding model configuration"""
    try:
        from model import get_embeddings
        from config import EMBEDDING_MODEL
        
        embeddings = get_embeddings()
        test_text = "Test embedding generation"
        vector = embeddings.embed_query(test_text)
        
        detail = f"model={EMBEDDING_MODEL} | dim={len(vector)}"
        return True, detail
    except Exception as exc:
        return False, f"{type(exc).__name__}: {exc}"


CHECKS: list[CheckSpec] = [
    # Core modules
    CheckSpec("OCR utilities", "ocr.py", "DoclingOCREngine", check_ocr_module),
    CheckSpec("Image helpers", "image.py", "SUPPORTED_FORMATS", check_image_helpers),
    CheckSpec("Audio metadata", "audio.py", "extract_file_metadata", check_audio_metadata),
    CheckSpec("Vector store", "vector_store.py", "create_vector_store", check_vector_store),
    CheckSpec("Text splitting", "langchain_text_splitters", "RecursiveCharacterTextSplitter", check_splitter),
    
    # Dependencies
    CheckSpec("Docling import", "docling", "DocumentConverter", check_docling_import),
    CheckSpec("Office fallbacks", "python-docx/pptx/openpyxl/pypdf", "import checks", check_office_fallbacks),
    CheckSpec("BM25 library", "rank_bm25", "BM25Okapi", check_bm25_import),
    
    # Connections & API keys
    CheckSpec("Ollama connectivity", "ollama", "list", check_ollama_connection),
    CheckSpec("Gemini API key", ".env", "GOOGLE_API_KEY", check_env_config),
    CheckSpec("Groq API key", ".env", "GROQ_API_KEY", check_groq_api_key),
    
    # RAG enhancements
    CheckSpec("Hybrid search (BM25+Vector)", "vector_store.py", "BM25Index", check_hybrid_search),
    CheckSpec("Query expansion", "vector_store.py", "expand_query", check_query_expansion),
    CheckSpec("Reranking", "vector_store.py", "rerank_documents", check_reranking),
    
    # LLM & Embeddings
    CheckSpec("LLM cascade config", "llm.py", "get_llm", check_llm_cascade),
    CheckSpec("Embeddings model", "model.py", "get_embeddings", check_embeddings),
    CheckSpec("LLM invocation", "llm.py", "invoke", check_llm_invocation),
]


def run_all_checks() -> bool:
    print("\n" + "=" * 72)
    print("SCOPED RAG - SMOKE TEST REPORT")
    print("=" * 72)

    results = []
    for spec in CHECKS:
        try:
            ok, detail = spec.runner()
        except Exception as exc:  # Defensive fallback
            ok, detail = False, f"Unexpected error: {type(exc).__name__}: {exc}"

        status = "PASS" if ok else "FAIL"
        print(f"[{status}] {spec.file} :: {spec.target} :: {spec.name}")
        print(f"    {detail}")
        results.append(ok)

    print("\n" + "=" * 72)
    print("SUMMARY")
    print("=" * 72)
    passed = sum(1 for r in results if r)
    total = len(results)
    print(f"Passed {passed}/{total} checks")

    failing = [spec for spec, ok in zip(CHECKS, results) if not ok]
    if failing:
        print("\nFailing checks:")
        for spec in failing:
            print(f"  - {spec.file} :: {spec.target} ({spec.name})")

    return passed == total


if __name__ == "__main__":
    sys.exit(0 if run_all_checks() else 1)
